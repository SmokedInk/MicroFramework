# -*- coding: utf-8 -*-
# @Author:  SmokedInk
# @Title:   IOWord
# @Time:    2019-11-15 13:56:05

from docx import Document
from win32com import client as wc


def write_type_info_docx(file_path):
    """
    示例：写入docx文件带有格式的内容
    :param file_path: 文件路径
    :return: None
    """
    # 获取docx文件对象
    document = Document()
    # 添加标题，并设置级别，范围：0 至 9，默认为1
    document.add_heading('Document Title', 0)
    # 添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
    p = document.add_paragraph('A plain paragraph having some ')
    # 在段落后面追加文本，并可设置样式
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    # 添加项目列表（前面一个小圆点）
    document.add_paragraph('first item in unordered list', style='List Bullet')
    document.add_paragraph('second item in unordered list', style='List Bullet')

    # 添加项目列表（前面数字）
    document.add_paragraph('first item in ordered list', style='List Number')
    document.add_paragraph('second item in ordered list', style='List Number')

    # 添加图片
    # document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    # 添加表格：一行三列
    # 表格样式参数可选：
    # Normal Table
    # Table Grid
    # Light Shading、 Light Shading Accent 1 至 Light Shading Accent 6
    # Light List、Light List Accent 1 至 Light List Accent 6
    # Light Grid、Light Grid Accent 1 至 Light Grid Accent 6
    # 太多了其它省略...
    # table = document.add_table(rows=1, cols=3, style='Light Shading Accent 2')
    # # 获取第一行的单元格列表
    # hdr_cells = table.rows[0].cells
    # # 下面三行设置上面第一行的三个单元格的文本值
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     # 表格添加行，并返回行所在的单元格列表
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc
    #
    # document.add_page_break()

    # 保存.docx文档
    document.save(file_path)


def write_docx(file_path, data):
    """
    写入docx文件内容
    :param file_path: 文件路径
    :param data: 待写入文件内容
    :return: None
    """
    # 获取docx文件对象
    document = Document()
    title = data.get("title")
    paragraphs = data.get("paragraphs")
    # 添加标题，并设置级别，范围：0 至 9，默认为1
    document.add_heading(title, 0)
    # 添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    # 在段落后面追加文本，并可设置样式
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    # 保存文件
    document.save(file_path)
    print("写入成功")


def written_docx(file_path, data):
    """
    追加写入docx文件内容
    :param file_path: 文件路径
    :param data: 待写入文件内容
    :return: None
    """
    document = Document(file_path)
    for paragraph in data:
        if type(paragraph) != tuple:
            paragraph = (paragraph, "Normal")
        print(paragraph)
        document.add_paragraph(paragraph[0], paragraph[1])
    # 保存文件
    document.save(file_path)
    print("插入成功")
    pass


def read_docx(file_path):
    """
    读取docx文档内容
    :param file_path: 文件路径
    :return: 文件内容
    """
    # 获取文档对象
    document = Document(file_path)
    print("段落数:" + str(len(document.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    text = []
    for para in document.paragraphs:
        text.append((para.text, para.style.name))

    # 表格的读取
    # tbs = file.tables
    # for tb in tbs:
    #     # 行
    #     for row in tb.rows:
    #         # 列
    #         for cell in row.cells:
    #             print(cell.text)

    return text


def doc_info_to_docx(file_path):
    """
    将doc文件更改为docx文件
    :param file_path: doc文件路径
    :return: docx文件路径
    """
    word = wc.Dispatch("Word.Application")
    file_dir = file_path.split("/" and "\\")
    file_dir_path = "\\".join(file_dir[:-1])
    file_name = file_dir[-1].split(".doc")[0]
    doc_file_path = file_dir_path + "\\" + file_name + ".doc"
    docx_file_path = file_dir_path + "\\" + file_name + ".docx"
    doc = word.Documents.Open(doc_file_path)
    doc.SaveAs(docx_file_path, 12)     # 12为docx
    doc.Close()
    word.Quit()
    return docx_file_path


if __name__ == '__main__':
    file_path = r"..\data\test.docx"
    # data1 = {
    #     "title": "This is the title!",
    #     "paragraphs": [
    #         "This is first paragraph",
    #         "This is second paragraph",
    #         "A plain paragraph having some"
    #     ]
    # }
    # data2 = ["A plain paragraph"]
    # write_docx(file_path, data1)
    # written_docx(file_path, data2)
    # doc_info_to_docx(file_path)
    for i in read_docx(file_path):
        print(i)
